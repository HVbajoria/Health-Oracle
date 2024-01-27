import streamlit as st
import time
from azure.cognitiveservices.vision.customvision.prediction import CustomVisionPredictionClient
from msrest.authentication import ApiKeyCredentials
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

# Replace with your endpoint and prediction key
ENDPOINT = "https://tuberculosisai-prediction.cognitiveservices.azure.com/"
PREDICTION_KEY = "eaf71df678cd4ddcb1c7d29ce1778260"
API_TOKEN = 'HJaeFtURjyS_Li39VxQJokDxSS4QUpVz4thesMoF'
account_id ="7b9a0aec2d148d684d28a3c7645ce0a7"

# Create a prediction client
credentials = ApiKeyCredentials(in_headers={"Prediction-key": PREDICTION_KEY})
predictor = CustomVisionPredictionClient(ENDPOINT, credentials)

st.set_page_config(page_title="HealthOracle: Decode Your Health")

doctors = [
    {
        "name": "Dr. Harshavardhan Bajoria",
        "specialization": "Pulmonologist",
        "location": "Kolkata",
        "available_days": "Mon, Tue, Fri",
        "contact": "hvbajoria@hotmail.com",
    },
    {
        "name": "Dr. Soumya Upadhyay",
        "specialization": "Chest Physician",
        "location": "Mumbai",
        "available_days": "Wed, Thu, Sat",
        "contact": "usoumya19@gmail.com",
    },
    # Add more doctors here...
]

def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a friendly medical assistant" },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/meta/llama-2-7b-chat-int8", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("chatbot_conversation.docx")

def runner():

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")

    if user_input:
        if user_input:
            bot_reply = bot_response(user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on end button to remove chat", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("chatbot_conversation.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            )

causes = """
Tuberculosis, caused by Mycobacterium tuberculosis, spreads through airborne droplets released during activities like speaking or coughing. Transmission is more likely indoors and in crowded settings. Latent TB is not contagious, and those on active TB treatment become non-contagious after 2-3 weeks. Drug-resistant TB results from genetic changes and issues like incorrect drug use, inadequate prescriptions, drug availability, poor quality, or absorption problems, contributing to the development and spread of resistant strains.
"""
symptoms = """
Tuberculosis (TB) infection progresses through primary infection, latent TB infection, and active TB disease stages. Primary infection is often asymptomatic, with flu-like symptoms if present. Latent TB infection is asymptomatic, as the immune system contains the germs. Active TB disease occurs when the immune system can't control the infection, manifesting with gradual symptoms in the lungs or other body parts. Extralung TB symptoms vary depending on the infected area. In children, symptoms differ by age group, ranging from typical adult-like symptoms in teenagers to distinctive signs in infants, such as growth issues and neurological symptoms.
"""
treat = """
TB treatment spans 3-9 months; adherence is vital for success. DOT or self-administration may be employed. Common drugs include Isoniazid, Rifampin, and more. Regular check-ups ensure progress and monitor side effects. If symptoms arise, promptly consult your healthcare provider.
"""

st.markdown(
    """
        <style>
            [data-testid="stSidebarNav"] {
                background-repeat: no-repeat;                
            }
            [data-testid="stSidebarNav"]::before {
                content: "Health Oracle";
                margin-left: 20px;
                margin-top: 20px;

                font-size: 30px;
                text-align: center;
                position: relative;
            }
        </style>
        """,
    unsafe_allow_html=True,
)
st.title("HealthOracle")
st.text(
    "Upload an image of a close up of a CT scan and we will tell you whether you have tuberculosis or not."
)
# read images.zip as a binary file and put it into the button
with open("test.zip", "rb") as fp:
    btn = st.download_button(
        label="Download test images",
        data=fp,
        file_name="test.zip",
        mime="application/zip",
    )
image = st.file_uploader(
    "Upload Image", type=["jpg", "jpeg", "png", "webp"], accept_multiple_files=False
)

if image is not None:
    disp = False
    
    with image:
        st.image(image, caption="Your CT Scan", width=350)
        image_data = image.read()
        results = predictor.classify_image("2ca27c6a-33bc-45db-8955-f47e4071fd1e", "Iteration1", image_data)
    disp = True
    
    c = st.image("loader.gif")
    time.sleep(3)
    c.empty()
    # Process and display the results
    if results.predictions:
        st.subheader("Prediction Results:")
        name="unknown"
        predict=0
        for prediction in results.predictions:
            if prediction.probability > predict and prediction.probability > 0.5:
                predict = prediction.probability
                name = prediction.tag_name

    if name!="unknown":
        st.success(f"Detected {name} with high confidence", icon='📃')
        if name == "Tuberculosis":
            st.write(
                """
                Tuberculosis (TB) is a serious illness that mainly affects the lungs. The germs that cause tuberculosis are a type of bacteria.
                Tuberculosis can spread when a person with the illness coughs, sneezes or sings. This can put tiny droplets with the germs into the air. Another person can then breathe in the droplets, and the germs enter the lungs.
                """
            )
            st.image("images/tuberculosis.png", caption="Glioma", width=350)
            st.write("More Info")

            tab1, tab2, tab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with tab1:
                st.write(causes)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )
            with tab2:
                st.write(symptoms)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )
            with tab3:
                st.write(treat)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment="Tuberculosis"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')

    else:
        st.text("No disease detected")
    
