import streamlit as st
import time
from azure.cognitiveservices.vision.customvision.prediction import CustomVisionPredictionClient
from msrest.authentication import ApiKeyCredentials
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

database_endpoint="https://hvbajoria101.kintone.com/k/v1/record.json?"
database_headers={'X-Cybozu-API-Token':'LeX70V7wU3KdgKN6JkzOlLkLK8nShxoEFbuF1ZWj', 'Content-Type': 'application/json'}

doctors = []

# Fetching doctors from database
for i in range(3,8):
    database_data = {
        'app':1,
        'id':i
    }

    database_response = requests.get(f"{database_endpoint}", headers=database_headers, json=database_data)
    
    doctor={}
    doctor["name"]=database_response.json()["record"]["Text"]["value"]
    doctor["specialization"]=database_response.json()["record"]["Text_0"]["value"]
    doctor["location"]=database_response.json()["record"]["Text_1"]["value"]
    doctor["available_days"]=database_response.json()["record"]["Text_3"]["value"]
    doctor["contact"]=database_response.json()["record"]["Text_2"]["value"]
    doctors.append(doctor)

# Replace with your endpoint and prediction key
ENDPOINT = "https://centralindia.api.cognitive.microsoft.com/"
PREDICTION_KEY = "cbc6c04b3e7c467eb5e7a294e9535afb"
API_TOKEN = 'HJaeFtURjyS_Li39VxQJokDxSS4QUpVz4thesMoF'
account_id ="7b9a0aec2d148d684d28a3c7645ce0a7"

# Create a prediction client
credentials = ApiKeyCredentials(in_headers={"Prediction-key": PREDICTION_KEY})
predictor = CustomVisionPredictionClient(ENDPOINT, credentials)

st.set_page_config(page_title="HealthOracle: Decode Your Health")

def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a friendly medical assistant" },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/meta/llama-2-7b-chat-int8", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("HealthOracle_Chat.docx")

def runner():

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")

    if user_input:
        if user_input:
            bot_reply = bot_response(user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on end button to remove chat", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("HealthOracle_Chat.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            )

gcauses = """
The exact causes of glioma, a type of brain tumor, are not fully understood. However, certain risk factors have been identified. These include exposure to ionizing radiation, a family history of glioma, and certain genetic disorders such as neurofibromatosis type 1 and Li-Fraumeni syndrome. While these factors may increase the risk, in many cases, the underlying cause of glioma remains unknown.
"""
gsymptoms = """
Gliomas can have significant symptoms on brain function and overall health. As the tumor grows, it can exert pressure on surrounding brain tissue, leading to symptoms such as headaches, seizures, difficulty speaking or understanding language, memory problems, changes in personality or mood, and neurological deficits like weakness or loss of sensation in the limbs. The severity and specific symptoms experienced by an individual can vary depending on the location, size, and grade of the glioma.
"""
gtreat = """
The treatment of glioma depends on several factors, including the tumor's location, size, grade, and the patient's overall health. Treatment options may include surgery to remove the tumor, radiation therapy to target and kill cancer cells, and chemotherapy to destroy or slow down tumor growth. In some cases, a combination of these treatments may be used. The choice of treatment is determined by a multidisciplinary team of medical professionals and is tailored to the individual patient's needs and circumstances. Regular monitoring and follow-up care are essential to assess the tumor's response to treatment and manage any potential side symptoms.
"""

mcauses = """The exact causes of meningioma, a type of brain tumor, are not well understood. However, certain risk factors have been identified, including radiation exposure, such as previous radiation therapy to the head, and certain genetic conditions like neurofibromatosis type 2 (NF2). Hormonal factors, such as increased levels of estrogen, have also been associated with an increased risk of developing meningiomas. Nonetheless, the underlying cause of most meningiomas remains unknown.
"""
msymptoms = """Meningiomas can have varying symptoms depending on their size, location, and growth rate. Some meningiomas may not cause noticeable symptoms and can be incidentally discovered during imaging tests conducted for unrelated reasons. However, when symptoms do occur, they can include headaches, seizures, changes in vision or hearing, weakness or numbness in the limbs, and cognitive or personality changes. The specific symptoms and their severity can differ from person to person.
"""
mtreat = """
The treatment of meningioma depends on factors such as tumor size, location, and growth rate, as well as the individual's overall health. Treatment options may include observation with regular monitoring for slow-growing or asymptomatic tumors, surgery to remove the tumor, radiation therapy to target and destroy cancer cells, and in some cases, medication to manage symptoms or slow down tumor growth. The choice of treatment is based on a thorough evaluation by a multidisciplinary team of healthcare professionals and is tailored to the specific needs of each patient. Regular follow-up care is important to assess the tumor's response to treatment and address any potential complications or recurrence.
"""

pcauses = """
The exact causes of pituitary tumors, also known as pituitary adenomas, are not fully understood. However, certain factors may increase the risk of their development. These include genetic conditions like multiple endocrine neoplasia type 1 (MEN1) and Carney complex, as well as rare hereditary syndromes such as familial isolated pituitary adenoma. Hormonal imbalances, exposure to certain chemicals, and head injuries have also been suggested as potential contributing factors. However, in many cases, the underlying cause of pituitary tumors remains unknown."""
psymptoms = """Pituitary tumors can have diverse symptoms depending on their size, location, and hormone production. They can disrupt the normal functioning of the pituitary gland, leading to hormonal imbalances and associated symptoms. The specific symptoms can vary widely, ranging from vision problems and headaches due to pressure on nearby structures, to hormonal disturbances resulting in issues such as infertility, growth abnormalities, changes in body composition, and metabolic problems. The symptoms of pituitary tumors are highly dependent on the specific hormones involved and the individual's overall health.
"""
ptreat = """The treatment of pituitary tumors depends on several factors, including the tumor's size, hormone production, and the individual's overall health. Treatment options may include medication to regulate hormone levels, surgery to remove the tumor, radiation therapy to destroy tumor cells, or a combination of these approaches. The choice of treatment is determined by a multidisciplinary team of medical professionals and is tailored to the individual patient's needs and circumstances. Regular monitoring and follow-up care are often necessary to manage hormone levels, monitor tumor growth, and ensure optimal treatment outcomes.
"""
st.markdown(
    """
        <style>
            [data-testid="stSidebarNav"] {
                background-repeat: no-repeat;                
            }
            [data-testid="stSidebarNav"]::before {
                content: "Health Oracle";
                margin-left: 20px;
                margin-top: 20px;

                font-size: 30px;
                text-align: center;
                position: relative;
            }
        </style>
        """,
    unsafe_allow_html=True,
)

def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "HealthOracle: Lung Lens"
  
# left_co, cent_co,last_co = st.columns(3)
# with cent_co:
#     st.image("images/logo.png", width=200)

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
st.text(
    "Upload an image of a close up of a tumerous MRI scan and we will tell you what type it is."
)
# read images.zip as a binary file and put it into the button
with open("test.zip", "rb") as fp:
    btn = st.download_button(
        label="Download test images",
        data=fp,
        file_name="test.zip",
        mime="application/zip",
    )
image = st.file_uploader(
    "Upload Image", type=["jpg", "jpeg", "png", "webp"], accept_multiple_files=False
)

if image is not None:
    disp = False
    
    with image:
        st.image(image, caption="Your MRI Scan", width=350)
        image_data = image.read()
        results = predictor.classify_image("e9f1afe5-1690-4ce7-88a4-23a05ee90732", "Iteration1", image_data)
    disp = True
    
    c = st.image("loader.gif")
    time.sleep(3)
    c.empty()
    # Process and display the results
    if results.predictions:
        st.subheader("Prediction Results:")
        name="unknown"
        predict=0
        for prediction in results.predictions:
            if prediction.probability > predict and prediction.probability > 0.5:
                predict = prediction.probability
                name = prediction.tag_name

    if name!="unknown":
        st.success(f"Detected {name} with high confidence", icon='📃')
        if name == "adenocarcinoma":
            st.write(
                """
                Glioma is a brain tumor that develops from glial cells. Its exact causes are not fully known, but risk factors include radiation exposure and certain genetic disorders. Gliomas can affect brain function, causing headaches, seizures, and neurological deficits. MRI is used to detect and evaluate gliomas, showing abnormal masses or areas of increased signal intensity. The size, location, and enhancement pattern of the tumor help determine its grade and guide treatment decisions.
                """
            )
            st.image("images/glioma.webp", caption="Glioma", width=350)
            st.write("More Info")

            tab1, tab2, tab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with tab1:
                st.write(gcauses)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )
            with tab2:
                st.write(gsymptoms)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )
            with tab3:
                st.write(gtreat)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/glioma/symptoms-causes/syc-20350251)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')

            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif (
            name == "large cell carcinoma"
        ):
            st.write(
                """
                Meningioma is a brain tumor that originates from the meninges, the protective membranes covering the brain and spinal cord. Its exact cause is unknown, but risk factors include radiation exposure, certain genetic conditions, and hormonal factors. Meningiomas can vary in symptoms depending on size and location. MRI is commonly used to detect and evaluate meningiomas, showing well-defined masses with a dural tail.
                """
            )
            st.image("images/Meningioma.jfif", caption="Meningioma", width=350)
            st.write("Known Carried Diseases")
            btab1, btab2, btab3 = st.tabs(
                ["Causes", "symptoms", "Treatment"]
            )
            with btab1:
                st.write(mcauses)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.cancer.gov/rare-brain-spine-tumor/tumors/meningioma)"
                )
            with btab2:
                st.write(msymptoms)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.cancer.gov/rare-brain-spine-tumor/tumors/meningioma)"
                )
            with btab3:
                st.write(mtreat)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.cancer.gov/rare-brain-spine-tumor/tumors/meningioma)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif name == "squamous cell carcinoma":
            st.write(
                """
                A pituitary tumor, also known as pituitary adenoma, is a non-cancerous growth in the pituitary gland. It can be functioning or non-functioning, causing hormonal imbalances or symptoms due to its size. Symptoms may include headaches, vision problems, fatigue, and hormonal disturbances. Diagnosis involves imaging tests like MRI, and treatment options include medication, surgery, or radiation therapy.
                """
            )
            st.image("images/petu.jfif", caption="Pituitary", width=350)
            st.write("Known Carried Diseases")
            ctab1, ctab2, ctab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with ctab1:
                st.write(pcauses)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/pituitary-tumors/symptoms-causes/syc-20350548)"
                )
            with ctab2:
                st.write(psymptoms)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/pituitary-tumors/symptoms-causes/syc-20350548)"
                )
            with ctab3:
                st.write(ptreat)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/pituitary-tumors/symptoms-causes/syc-20350548)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

    else:
        st.text("No disease detected")
    
